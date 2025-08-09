# bot.py
import os
import re
import sys
import tempfile
import shutil
import subprocess
import unicodedata
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor

from dotenv import load_dotenv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder, MessageHandler, filters, CommandHandler,
    ContextTypes, CallbackQueryHandler
)

# OCR / PDF
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import numpy as np
import cv2  # OpenCV
from PyPDF2 import PdfMerger, PdfReader

# DOCX
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# ========= تحميل التوكن =========
load_dotenv()
def get_token() -> str:
    token = os.getenv("TELEGRAM_BOT_TOKEN")
    if not token or not re.match(r"^\d+:[A-Za-z0-9_-]{20,}$", token):
        print("❌ TELEGRAM_BOT_TOKEN مفقود/غير صحيح. ضع التوكن الصحيح في .env")
        sys.exit(1)
    return token
TOKEN = get_token()

# ========= إعدادات عامة =========
OCR_LANG = "ara+tur+eng"
OCR_DPI  = 500
WORKERS  = max(1, (os.cpu_count() or 2) - 1)

# (لو ويندوز) فعّل السطر وحدد المسار إذا احتجت:
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ========= نص: تطبيع/تصحيح + RTL =========
_AR_REPLACEMENTS = {
    "\u06BE": "\u0647", "\u06C1": "\u0647", "\u06D5": "\u0629",
    "\u0649": "\u064A", "\u06CC": "\u064A", "\u06A9": "\u0643",
    "\u06AF": "\u063A", "\u0629\u0640": "\u0629",
}
def harmonize_ar_chars(t: str) -> str:
    return "".join(_AR_REPLACEMENTS.get(ch, ch) for ch in t)
def normalize_arabic_text(t: str) -> str:
    if not t: return ""
    t = unicodedata.normalize("NFKC", t)
    t = t.replace("\u0640", "")
    t = re.sub(r"[\u200e\u200f\u202a-\u202e]", "", t)
    t = harmonize_ar_chars(t)
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"[=~_*]{2,}", " ", t)
    return t.strip()

_AR_RUN = re.compile(r"[\u0600-\u06FF]+")
_MIRROR = str.maketrans({"(": ")", ")": "(", "[": "]", "]": "[", "{": "}", "}": "{", "<": ">", ">": "<", "«": "»", "»": "«"})
def _reverse_ar_run(run: str) -> str:
    return run[::-1].translate(_MIRROR)
def flip_arabic_runs(s: str) -> str:
    return _AR_RUN.sub(lambda m: _reverse_ar_run(m.group(0)), s)
def looks_visual_arabic(text: str) -> bool:
    toks = _AR_RUN.findall(text)
    if len(toks) < 5: return False
    start_al = sum(t.startswith("ال") for t in toks)
    end_la   = sum(t.endswith("لا") for t in toks)
    return end_la > max(5, int(start_al * 1.5))
def fix_visual_arabic_if_needed(text: str) -> str:
    return flip_arabic_runs(text) if text and looks_visual_arabic(text) else text
def postprocess_text(text: str) -> str:
    text = normalize_arabic_text(text)
    text = fix_visual_arabic_if_needed(text)
    return text

# ========= محارف اتجاه ثنائية لملف TXT =========
RLM = "\u200f"   # Right-to-Left Mark
RLE = "\u202b"   # Right-to-Left Embedding
PDF = "\u202c"   # Pop Directional Formatting
ALM = "\u061C"   # Arabic Letter Mark (يضبط اتجاه الأرقام/اللاتيني قرب العربي)
# نطاق أوسع للأحرف العربية:
_AR_CHARS_WIDE = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]")
# أرقام ولاتيني:
_DIGITS = re.compile(r"(\d+)")
_LATIN  = re.compile(r"([A-Za-z@#_./:\\-]+)")

def wrap_line_rtl_if_ar(line: str) -> str:
    """لفّ السطر العربي بمحارف اتجاه قوية + حماية للأرقام/اللاتيني بـ ALM."""
    if not line or not _AR_CHARS_WIDE.search(line):
        return line

    # أحط ALM قبل/بعد الأرقام حتى تبقى ضمن سياق RTL
    line = _DIGITS.sub(lambda m: f"{ALM}{m.group(1)}{ALM}", line)
    # وأيضًا قبل/بعد المقاطع اللاتينية/الروابط/الإيميلات
    line = _LATIN.sub(lambda m: f"{ALM}{m.group(1)}{ALM}", line)

    # لفّ السطر كله في تضمين RTL + علامات RLM لتثبيت الاتجاه عند الأطراف
    return f"{RLE}{RLM}{line}{RLM}{PDF}"

# ========= سكور جودة للنص =========
_AR_LETTERS = re.compile(r"[؀-ۿ]+")
_TR_LATIN   = re.compile(r"[çğıöşüÇĞİÖŞÜ]")
def score_text_quality(txt: str) -> float:
    if not txt: return 0.0
    total = len(txt)
    ar = len("".join(_AR_LETTERS.findall(txt)))
    tr = len("".join(_TR_LATIN.findall(txt)))
    words = [w for w in re.findall(r"\w+", txt) if len(w) >= 2]
    avgw  = (sum(len(w) for w in words) / max(1, len(words))) if words else 0
    noise = len(re.findall(r"[=~_*]{2,}", txt)) + len(re.findall(r"\d{5,}", ""))
    return (ar*2 + tr*1.5) / max(1,total) + (avgw/20) - (noise*0.05)

# ========= تحسين الصورة + deskew =========
def render_page_image(page, dpi: int = OCR_DPI) -> Image.Image:
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

def pil_to_cv(img: Image.Image):
    return cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)

def deskew(gray: np.ndarray) -> np.ndarray:
    coords = np.column_stack(np.where(gray < 200))
    if coords.size == 0: return gray
    angle = cv2.minAreaRect(coords)[-1]
    angle = -(90 + angle) if angle < -45 else -angle
    (h, w) = gray.shape[:2]
    M = cv2.getRotationMatrix2D((w//2, h//2), angle, 1.0)
    return cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

def preprocess_for_ocr(img: Image.Image) -> Image.Image:
    cv = pil_to_cv(img)
    gray = cv2.cvtColor(cv, cv2.COLOR_BGR2GRAY)
    gray = cv2.bilateralFilter(gray, 9, 75, 75)
    gray = deskew(gray)
    thr = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                cv2.THRESH_BINARY, 31, 10)
    kernel = np.ones((1,1), np.uint8)
    thr = cv2.morphologyEx(thr, cv2.MORPH_OPEN, kernel, iterations=1)
    thr = cv2.morphologyEx(thr, cv2.MORPH_CLOSE, kernel, iterations=1)
    return Image.fromarray(cv2.cvtColor(thr, cv2.COLOR_GRAY2RGB))

# ========= OCR =========
def tesseract_try(img: Image.Image, lang: str) -> str:
    base = '-c preserve_interword_spaces=1 -c load_system_dawg=0 -c load_freq_dawg=0 -c tessedit_char_blacklist=ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    configs = [
        rf'--oem 1 --psm 6 {base}',
        rf'--oem 1 --psm 4 {base}',
        rf'--oem 1 --psm 3 {base}',
        rf'--oem 3 --psm 6 {base}',
        rf'--oem 1 --psm 13 {base}',
    ]
    best_txt, best_sc = "", -1e9
    for cfg in configs:
        try:
            raw = pytesseract.image_to_string(img, lang=lang, config=cfg)
            txt = postprocess_text(raw)
            sc = score_text_quality(txt)
            if sc > best_sc:
                best_sc, best_txt = sc, txt
        except Exception:
            continue
    return best_txt.strip()

def ocr_pages_to_list(prep_images):
    with ThreadPoolExecutor(max_workers=WORKERS) as pool:
        results = list(pool.map(lambda im: tesseract_try(im, OCR_LANG), prep_images))
    return [postprocess_text(t) for t in results]

# ========= Searchable PDF =========
def has_cmd(cmd: str) -> bool:
    return shutil.which(cmd) is not None

def searchable_pdf_with_ocrmypdf(input_pdf: str, output_pdf: str) -> bool:
    if not has_cmd("ocrmypdf"):
        return False
    try:
        cmd = [
            "ocrmypdf",
            "-l", OCR_LANG,
            "--force-ocr", "--rotate-pages", "--deskew", "--clean-final",
            "--optimize", "0",
            "--jobs", str(WORKERS),
            input_pdf, output_pdf
        ]
        subprocess.run(cmd, check=True)
        return True
    except Exception:
        return False

def searchable_pdf_with_tesseract_only_from_images(orig_images, output_pdf: str) -> bool:
    try:
        merger = PdfMerger()
        for img in orig_images:
            pdf_bytes = pytesseract.image_to_pdf_or_hocr(img, lang=OCR_LANG, extension='pdf')
            tmp_page = BytesIO(pdf_bytes)
            merger.append(PdfReader(tmp_page))
        with open(output_pdf, "wb") as f:
            merger.write(f)
        return True
    except Exception:
        return False

# ========= أدوات صور داخلية + DOCX =========
def _set_rtl_para(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), "1"); pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'), "1"); rPr.append(rtl)

def _page_blocks(page):
    d = page.get_text("rawdict")
    blocks = d.get("blocks", []) if isinstance(d, dict) else []
    blocks.sort(key=lambda b: (b.get("bbox", [0,0,0,0])[1], b.get("bbox", [0,0,0,0])[0]))
    return blocks

def _render_page_for_cropping(page, dpi=240):
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    pil = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    scale_x = pix.width / page.rect.width
    scale_y = pix.height / page.rect.height
    return pil, scale_x, scale_y

def _embedded_text_from_rawdict(page) -> str:
    try:
        d = page.get_text("rawdict")
        lines = []
        for b in d.get("blocks", []):
            if b.get("type", 0) != 0:  # نص فقط
                continue
            for l in b.get("lines", []):
                spans = [s.get("text","") for s in l.get("spans", [])]
                line = " ".join(spans).strip()
                if line:
                    lines.append(line)
        return "\n".join(lines).strip()
    except Exception:
        return ""

def make_docx_text_only(ocr_text_by_page, out_path):
    doc = Document()
    for i, raw in enumerate(ocr_text_by_page, 1):
        if i > 1:
            doc.add_page_break()
        txt = postprocess_text(raw or "")
        for para in txt.split("\n"):
            p = doc.add_paragraph(para)
            _set_rtl_para(p)
    doc.save(out_path)

def make_docx_with_inline_images(pdf_path, ocr_text_by_page, out_path):
    """
    يدرج الصور الداخلية في مواقعها التقريبية وبحجمها النسبي،
    ويضيف نص الصفحة (OCR أولاً، ثم نص مضمّن fallback، ثم OCR فوري للصفحة عند الحاجة).
    """
    doc = Document()
    # عرض المحتوى داخل هوامش Word
    sec = doc.sections[0]
    page_width_in  = float(sec.page_width)  / 914400.0
    left_in        = float(sec.left_margin) / 914400.0
    right_in       = float(sec.right_margin)/ 914400.0
    content_w_in   = max(1.0, page_width_in - left_in - right_in)

    with fitz.open(pdf_path) as pdf:
        for pno in range(len(pdf)):
            page = pdf[pno]
            blocks = _page_blocks(page)
            page_img, sx, sy = _render_page_for_cropping(page, dpi=240)

            if pno > 0:
                doc.add_page_break()

            # 1) نص الصفحة: OCR -> embedded -> OCR سريع على الصفحة عند الحاجة
            txt = ocr_text_by_page[pno] if pno < len(ocr_text_by_page) else ""
            txt = txt or _embedded_text_from_rawdict(page)
            if not txt.strip():
                # OCR سريع fallback على صورة الصفحة مباشرةً
                quick = preprocess_for_ocr(render_page_image(page, dpi=OCR_DPI))
                txt = tesseract_try(quick, OCR_LANG)
            txt = postprocess_text(txt or "")

            # قسّم لفقرات
            paragraphs = [p for p in (t.strip() for t in txt.split("\n")) if p]
            if not paragraphs:
                paragraphs = [" "]  # slot واحد على الأقل لحقن الصور

            # 2) اجمع الصور ونِسَبها
            img_entries = []
            page_w = float(page.rect.width)
            page_h = float(page.rect.height)
            for b in blocks:
                if b.get("type", 0) != 1:
                    continue
                x0, y0, x1, y1 = b.get("bbox", [0,0,0,0])
                y_mid_ratio = min(0.999, max(0.0, ((y0 + y1) / 2.0) / page_h))
                w_ratio = max(0.0, (x1 - x0) / page_w)
                crop = page_img.crop((int(x0*sx), int(y0*sy), int(x1*sx), int(y1*sy)))
                if crop.width < 30 or crop.height < 30:
                    continue
                img_entries.append({
                    "ratio": y_mid_ratio,
                    "width_ratio": w_ratio,
                    "x_center_ratio": ((x0 + x1) / 2.0) / page_w,
                    "img": crop
                })

            # 3) حقن الصور وفق موضعها النسبي
            total_slots = len(paragraphs) + 1  # أماكن بين الفقرات
            injected_map = {}
            for ent in img_entries:
                idx = int(round(ent["ratio"] * total_slots))
                idx = min(max(idx, 0), total_slots - 1)
                injected_map.setdefault(idx, []).append(ent)

            for slot in range(total_slots):
                # صور في هذا الموضع
                for ent in injected_map.get(slot, []):
                    buf = BytesIO(); ent["img"].save(buf, format="PNG"); buf.seek(0)
                    width_in = max(0.8, content_w_in * max(0.1, min(ent["width_ratio"], 1.0)))
                    pic_para = doc.add_paragraph()
                    xc = ent["x_center_ratio"]
                    if xc < 0.33:
                        pic_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif xc > 0.66:
                        pic_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pic_para.add_run().add_picture(buf, width=Inches(width_in))
                    doc.add_paragraph("")

                # فقرة النص (إن وجدت)
                if slot < len(paragraphs):
                    p = doc.add_paragraph(paragraphs[slot])
                    _set_rtl_para(p)

    doc.save(out_path)

# ========= إرسال نص طويل (اختياري) =========
async def send_chunked_text(update: Update, text: str, header: str):
    text = postprocess_text(text)
    if not text:
        await update.message.reply_text("لم يتم العثور على نص.")
        return

    # لفّ السطور العربية قبل التجزئة ليُعرض RTL في تيليجرام/المحررات
    prepared_lines = []
    for line in text.split("\n"):
        prepared_lines.append(wrap_line_rtl_if_ar(line.strip()))
    text = "\n".join(prepared_lines)

    chunk_size = 3500
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    await update.message.reply_text(header + f"\n(أجزاء: {len(chunks)})")
    for idx, ch in enumerate(chunks, 1):
        await update.message.reply_text(f"جزء {idx}:\n{ch}")

# ========= رندر الصفحات (أصلية + محضّرة) =========
def render_pages_both(pdf_path: str):
    orig_imgs, prep_imgs = [], []
    with fitz.open(pdf_path) as doc:
        for page in doc:
            orig = render_page_image(page)
            prep = preprocess_for_ocr(orig)
            orig_imgs.append(orig); prep_imgs.append(prep)
    return orig_imgs, prep_imgs

# ========= المعالجة =========
async def build_and_send_output(update: Update, pdf_path: str, choice: str, context: ContextTypes.DEFAULT_TYPE):
    status_msg = await update.message.reply_text("⏳ جاري المعالجة…")
    temp_dir = tempfile.mkdtemp(prefix="tg_pdf_")
    try:
        out_file_path = None
        used_ocr = False

        # رندر الصفحات مرة واحدة
        orig_images, prep_images = render_pages_both(pdf_path)

        if choice == "SEARCHABLE_PDF":
            out_pdf = os.path.join(temp_dir, "searchable.pdf")
            ok = searchable_pdf_with_ocrmypdf(pdf_path, out_pdf)
            if not ok:
                ok = searchable_pdf_with_tesseract_only_from_images(orig_images, out_pdf)
            if not ok:
                await status_msg.edit_text("⚠️ تعذّر إنشاء PDF قابل للبحث. سأُرجع DOCX.")
                choice = "DOCX"
            else:
                used_ocr = True
                out_file_path = out_pdf

        if choice in ("TXT", "DOCX", "DOCX_INLINE"):
            await status_msg.edit_text("🔍 OCR عبر Tesseract …")
            per_page_text = ocr_pages_to_list(prep_images)
            used_ocr = True

            if choice == "TXT":
                out_file_path = os.path.join(temp_dir, "output.txt")

                lines_out = []
                for i, raw in enumerate(per_page_text, 1):
                    # عنوان الصفحة — لفّه أيضًا ليعرض RTL
                    lines_out.append(wrap_line_rtl_if_ar(f"--- صفحة {i} ---"))
                    txt = postprocess_text(raw or "")
                    for line in txt.split("\n"):
                        line = line.strip()
                        if not line:
                            lines_out.append("")
                            continue
                        # لفّ السطر العربي بمحارف اتجاه
                        line = wrap_line_rtl_if_ar(line)
                        lines_out.append(line)

                joined = "\n".join(lines_out)

                with open(out_file_path, "w", encoding="utf-8-sig") as f:
                    f.write(joined)

            elif choice == "DOCX":
                out_file_path = os.path.join(temp_dir, "output.docx")
                make_docx_text_only(per_page_text, out_path=out_file_path)
            else:  # DOCX_INLINE
                out_file_path = os.path.join(temp_dir, "output_inline.docx")
                make_docx_with_inline_images(pdf_path, per_page_text, out_path=out_file_path)

        caption = f"✅ تم الإنشاء ({'Searchable PDF' if choice=='SEARCHABLE_PDF' else choice})"
        if used_ocr:
            caption += " • Tesseract OCR"
        await update.message.reply_document(document=open(out_file_path, "rb"), caption=caption)
        await status_msg.delete()

    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

# ========= الواجهة =========
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 أرسل PDF (حتى لو صور). أستطيع تحويله إلى:\n"
        "• DOCX (نص فقط)\n"
        "• DOCX (نص + صور داخلية بمواقعها)\n"
        "• TXT\n"
        "• PDF قابل للبحث (يحافظ على الصور)\n"
        "أستخدم Tesseract (ara+tur+eng) وأصلّح اتجاه العربي."
    )

async def ask_output_format(update: Update, context: ContextTypes.DEFAULT_TYPE, pdf_path: str):
    context.user_data["last_pdf_path"] = pdf_path
    keyboard = [
        [InlineKeyboardButton("DOCX (نص فقط)", callback_data="OUT_DOCX"),
         InlineKeyboardButton("DOCX (نص + صور داخلية)", callback_data="OUT_DOCX_INLINE")],
        [InlineKeyboardButton("TXT",  callback_data="OUT_TXT")],
        [InlineKeyboardButton("PDF قابل للنسخ", callback_data="OUT_SEARCHABLE")],
    ]
    await update.message.reply_text("اختر صيغة الإخراج:", reply_markup=InlineKeyboardMarkup(keyboard))

async def handle_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    doc = update.message.document
    if not doc or not doc.file_name.lower().endswith(".pdf"):
        await update.message.reply_text("❌ الرجاء إرسال ملف PDF فقط.")
        return
    file_obj = await doc.get_file()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        await file_obj.download_to_drive(custom_path=tmp.name)
        saved_path = tmp.name
    await ask_output_format(update, context, saved_path)

async def on_choice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data
    pdf_path = context.user_data.get("last_pdf_path")
    if not pdf_path or not os.path.exists(pdf_path):
        await query.edit_message_text("⚠️ لا يوجد ملف PDF محفوظ. أرسل الملف من جديد.")
        return
    if data == "OUT_DOCX": choice = "DOCX"
    elif data == "OUT_DOCX_INLINE": choice = "DOCX_INLINE"
    elif data == "OUT_TXT": choice = "TXT"
    elif data == "OUT_SEARCHABLE": choice = "SEARCHABLE_PDF"
    else:
        await query.edit_message_text("❌ خيار غير معروف.")
        return
    await query.edit_message_text(f"جاري إنشاء: {choice} …")
    try:
        chat = await context.bot.get_chat(query.message.chat_id)
        dummy_msg = await chat.send_message("⏳ بدء المعالجة…")
        class DummyUpdate: message = dummy_msg
        await build_and_send_output(DummyUpdate(), pdf_path, choice, context)
    finally:
        try: os.remove(pdf_path)
        except Exception: pass

def main():
    app = ApplicationBuilder().token(TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Document.PDF, handle_pdf))
    app.add_handler(CallbackQueryHandler(on_choice))
    print("🤖 البوت شغّال…")
    app.run_polling()

if __name__ == "__main__":
    main()
